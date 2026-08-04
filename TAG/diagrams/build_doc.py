"""Companion review document for the TAG submission, styled on 'MS Fabric Review Notes.docx'."""
import os, copy
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(HERE, "out")
TAG = r"C:\Users\Shuo\OneDrive\文档\bcpensionshackathon\TAG"
TPL = os.path.join(TAG, "MS Fabric Review Notes.docx")
DST = os.path.join(TAG, "MS Fabric Platform Direction - Review Notes (DRAFT).docx")

TEAL = RGBColor(0x00, 0x77, 0x8A)
GREY = RGBColor(0x4D, 0x4D, 0x4F)
WIDTH = 6.5

doc = docx.Document(TPL)

# ---- clear the body but keep styles and the section properties -------------
body = doc.element.body
sectPr = body.find(qn('w:sectPr'))
for el in list(body):
    if el is not sectPr:
        body.remove(el)


def para(text="", style=None, size=None, bold=False, italic=False, color=None,
         space_after=6, align=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        r = p.add_run(text)
        if size:
            r.font.size = Pt(size)
        r.bold = bold
        r.italic = italic
        if color is not None:
            r.font.color.rgb = color
    return p


def h(text, level=1):
    return doc.add_paragraph(text, style=f"Heading {level}")


def bullets(items, style="List Paragraph", size=10.5):
    for it in items:
        p = doc.add_paragraph(style=style)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Inches(0.3)
        pPr = p._p.get_or_add_pPr()
        numPr = OxmlElement('w:numPr')
        for tag, val in (('w:ilvl', 0), ('w:numId', 1)):
            e = OxmlElement(tag); e.set(qn('w:val'), str(val)); numPr.append(e)
        pPr.append(numPr)
        if isinstance(it, tuple):
            r = p.add_run(it[0] + " — "); r.bold = True; r.font.size = Pt(size)
            r2 = p.add_run(it[1]); r2.font.size = Pt(size)
        else:
            p.add_run(it).font.size = Pt(size)


def table(headers, rows, widths=None, size=9.5, header_fill="00778A"):
    t = doc.add_table(rows=1, cols=len(headers))
    try:
        t.style = doc.styles["Info Table"]
    except KeyError:
        t.style = doc.styles["Normal Table"]
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    hdr = t.rows[0]
    for i, htxt in enumerate(headers):
        c = hdr.cells[i]
        c.text = ""
        p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(2)
        r = p.add_run(htxt); r.bold = True; r.font.size = Pt(size)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        sh = OxmlElement('w:shd'); sh.set(qn('w:fill'), header_fill)
        c._tc.get_or_add_tcPr().append(sh)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]; p.paragraph_format.space_after = Pt(2)
            r = p.add_run(str(val)); r.font.size = Pt(size)
            if i == 0 and len(headers) > 1:
                r.bold = True
    if widths:
        total = sum(widths)
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(WIDTH * w / total)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def figure(png, caption):
    doc.add_picture(png, width=Inches(WIDTH))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    para(caption, size=9, italic=True, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER,
         space_after=10)


# ============================================================ TITLE + METADATA
doc.add_paragraph("P0905 Microsoft Fabric — Platform Direction", style="Corp-Title")

meta = doc.add_table(rows=7, cols=2)
try:
    meta.style = doc.styles["Info Table"]
except KeyError:
    pass
for i, (k, v) in enumerate([
        ("Title", "Microsoft Fabric — Platform Direction (TAG submission companion)"),
        ("Date Created", "August 4, 2026"),
        ("Updated On", ""),
        ("Updated by", ""),
        ("Created By", ""),
        ("DEPARTMENT", "IT Engineering"),
        ("Status", "Draft — for TAG review")]):
    meta.rows[i].cells[0].text = ""
    r = meta.rows[i].cells[0].paragraphs[0].add_run(k); r.bold = True; r.font.size = Pt(9.5)
    meta.rows[i].cells[1].text = ""
    meta.rows[i].cells[1].paragraphs[0].add_run(v).font.size = Pt(9.5)
    meta.rows[i].cells[0].width = Inches(1.7)
    meta.rows[i].cells[1].width = Inches(4.8)
doc.add_paragraph()

# ==================================================================== PURPOSE
h("Purpose", 1)
para("This document accompanies the TAG submission deck for project P0905. It sets out the "
     "proposed Microsoft Fabric platform direction in enough detail for reviewers to assess it "
     "ahead of the session, and to record the decisions, assumptions and risks that TAG is being "
     "asked to endorse.", size=10.5)
para("The deck is deliberately high level. This document holds the supporting detail behind each "
     "of its diagrams. It is a companion to the existing MS Fabric Review Notes, which cover the "
     "Fabric product capabilities and Microsoft's published deployment patterns; the intent here "
     "is narrower — what BC Pension Corporation proposes to build, and how it will be secured, "
     "governed and operated.", size=10.5)
para("What TAG is being asked to approve: to proceed with a proof of concept of Microsoft Fabric "
     "on the architecture described below, run in isolated environments on masked or synthetic "
     "data, ending in a recommendation to leadership.", size=10.5, bold=True)

# =============================================================== ARCHITECTURE
h("High-level architecture", 1)
para("Data reaches Fabric only over private connectivity. Inside Fabric it lands in OneLake and "
     "moves through a medallion progression — Bronze for raw, append-only history, Silver for "
     "cleansed and quality-checked data, Gold for the modelled star schema that reporting reads. "
     "Orchestration is carried out by Fabric data pipelines driving parameter-driven Spark "
     "notebooks; near-real-time sources can instead be replicated with Fabric Mirroring.",
     size=10.5)
figure(os.path.join(OUT, "d1_architecture.png"),
       "Figure 1 — Logical architecture. Sources reach Fabric through the VNet or on-premises "
       "data gateway; platform services apply to every environment.")

h("What is distinctive about this design", 2)
bullets([
    ("Configuration, not code, per source", "sources, quality rules, models and security "
     "policies are declared as configuration held in git and promoted between environments. "
     "Onboarding a new source is a configuration change, which keeps the marginal effort per "
     "source small."),
    ("One trusted copy", "OneLake holds a single copy of the data; the SQL endpoint and semantic "
     "models read the same storage rather than taking further copies."),
    ("Security declared alongside the data", "row-level and column-level security and masking "
     "are part of the same configuration, applied per environment rather than hand-built."),
    ("Auditable by construction", "every source-to-Gold run is recorded, alongside Fabric's own "
     "lineage view."),
], size=10.5)

h("Layer responsibilities", 2)
table(["Layer", "Holds", "Access posture"],
      [["Bronze", "Raw landed data, append-only, retaining source history.",
        "Platform and engineering only. Never exposed to report consumers."],
       ["Silver", "Cleansed, de-duplicated and quality-checked data; failing rows quarantined.",
        "Engineering and stewards; masking applied as data lands."],
       ["Gold", "Modelled star schema — the governed data products.",
        "The layer reporting reads, subject to row and column security."],
       ["Serve", "Semantic models, Power BI reports and the SQL endpoint.",
        "Business consumers, through a published app."]],
      widths=[1.1, 3.0, 2.4])

# ================================================= WORKSPACE / ENVIRONMENTS
h("Workspace and environment model", 1)
para("Every environment is a hard boundary with its own workspaces, Entra groups, service "
     "principals, key vault and source connections. Production data never flows downward: lower "
     "environments run on masked or synthetic data. Promotion moves code and configuration only, "
     "executed by a deployment service principal.", size=10.5)
figure(os.path.join(OUT, "d2_environments.png"),
       "Figure 2 — Workspace, environment and deployment model.")

h("Isolation matrix", 2)
table(["Isolated per environment", "DEV", "UAT", "PROD"],
      [["Workspaces", "separate", "separate", "separate"],
       ["Capacity", "shared acceptable", "shared acceptable", "dedicated (recommended)"],
       ["Entra groups", "own", "own", "own"],
       ["Service principals", "own", "own", "own"],
       ["Key vault (secrets)", "own", "own", "own — production credentials live only here"],
       ["Source connection", "DEV source", "UAT source", "PROD source"],
       ["Data", "masked / synthetic", "masked / representative", "real (PII governed)"],
       ["Human change access", "build freely", "limited", "none standing — promotion only"]],
      widths=[2.1, 1.4, 1.4, 1.6])

h("Access matrix — who gets what, where", 2)
para("This is the core of an access-control review: which persona gets which level in which "
     "environment. The shape — broad in DEV, tight in PROD, promotion-only for changes — is the "
     "part to scrutinise.", size=10.5, italic=True)
table(["Persona / identity", "DEV", "UAT", "PROD"],
      [["Platform / Fabric admin", "Admin (elevated on request)", "Admin (elevated on request)",
        "Admin — just-in-time with approval; break-glass held separately"],
       ["Data engineer / developer", "Contributor or Member", "Viewer, or Contributor for fixes",
        "No standing access — changes via pipeline only"],
       ["Data steward", "Member", "Member", "Viewer plus governance tooling (no data edits)"],
       ["Report author / analyst", "Contributor", "Contributor (validate)",
        "Author only in a dedicated authoring workspace; Viewer on curated data"],
       ["Business consumer", "—", "Viewer (UAT sign-off)", "Viewer through a published app only"],
       ["Auditor / security", "Viewer + audit logs", "Viewer + audit logs", "Viewer + audit logs"],
       ["Deployment service principal", "Member (deploy)", "Member (deploy)",
        "Member — the only identity that changes production"],
       ["Runtime service principal", "Contributor (run)", "Contributor (run)",
        "Contributor (run)"],
       ["Gateway / connection account", "read-only on DEV source", "read-only on UAT source",
        "read-only on PROD source"]],
      widths=[1.8, 1.4, 1.5, 1.8], size=9)

h("Promotion", 2)
para("Changes flow DEV → UAT → PROD through Fabric deployment pipelines and the control-plane "
     "deployment scripts, executed by the deployment service principal. Only code and "
     "configuration promote — never data, and never connection secrets, which are held per "
     "environment in that environment's key vault.", size=10.5)

# ==================================================================== SECURITY
h("Security and access model", 1)
para("Access is granted to Microsoft Entra security groups, gated by Conditional Access and "
     "just-in-time elevation, and then narrows layer by layer inside Fabric — from tenant "
     "guardrails down to individual rows, columns and masked values. Automation runs as service "
     "principals with secrets in Key Vault, and every action is audited.", size=10.5)
figure(os.path.join(OUT, "d3_security.png"),
       "Figure 3 — Identity, defence in depth, and the governance operating model.")

h("Guiding principles", 2)
bullets([
    ("Group-based, never individual", "all access is granted to Entra security groups; adding or "
     "removing a person is a membership change, and the permission model itself never changes."),
    ("Least privilege and just-in-time", "privileged and production access is granted on demand "
     "with approval, not held standing."),
    ("People and machines do not mix", "humans never authenticate as service principals, and "
     "automation never runs as a personal account."),
    ("Production data stays in production", "lower environments use masked, synthetic or subset "
     "data."),
    ("Change production only by promotion", "no manual edits in production."),
    ("Private by default", "no public-internet path to any source."),
    ("Everything is auditable", "access grants, data reads and pipeline actions are all logged."),
], size=10.5)

h("Entra security groups", 2)
table(["Group (naming convention)", "Members", "Purpose"],
      [["BCPC-FAB-{ENV}-Engineers", "Data engineers", "Build access in that environment"],
       ["BCPC-FAB-{ENV}-ReportAuthors", "Analysts", "Author semantic models and reports"],
       ["BCPC-FAB-{ENV}-Consumers", "Business users", "View published reports through an app"],
       ["BCPC-DATA-{DOMAIN}-Stewards", "Data stewards", "Governance for a data domain"],
       ["BCPC-FAB-Platform-Admins", "Platform team",
        "Tenant and capacity administration — elevation required"],
       ["BCPC-FAB-Auditors", "Security and audit", "Read and audit-log access"],
       ["BCPC-FAB-SPN-{ENV}", "Service principals", "Automation identities; scopes API access"]],
      widths=[2.3, 1.5, 2.7])

h("Defence in depth", 2)
table(["Layer", "Control"],
      [["Tenant guardrails",
        "Service-principal scope, publish-to-web off, external sharing off, mandatory sensitivity "
        "labels, Conditional Access."],
       ["Capacity", "Production on dedicated capacity; workspace-to-capacity assignment "
        "controlled."],
       ["Workspace roles", "Least-privilege Admin / Member / Contributor / Viewer, assigned to "
        "groups and reviewed on a cadence."],
       ["Item permissions", "Per lakehouse, semantic model and report."],
       ["Data-level security", "OneLake row-level and column-level security plus dynamic and "
        "static masking, applied from configuration per environment."]],
      widths=[1.6, 4.9])

h("Machine identities", 2)
table(["Identity (per environment)", "Used for", "Access"],
      [["Deployment service principal", "CI/CD — provision and promote code and configuration",
        "Member on that environment's workspaces; the only identity that changes production"],
       ["Runtime service principal", "Scheduled pipeline and notebook execution",
        "Contributor on that environment's workspaces"],
       ["Gateway / connection account", "Source connectivity through the gateway",
        "Read-only on the source"],
       ["Managed identity (preferred where supported)", "Fabric-to-Azure resource access",
        "Scoped to the specific resource; no credential to manage"]],
      widths=[2.0, 2.2, 2.3])

h("Data protection", 2)
table(["Control", "Position"],
      [["PII minimisation", "Sensitive columns are masked or excluded as data lands in Silver, so "
        "sensitive values need never persist raw."],
       ["Row and column security", "Applied from configuration, consistently per environment."],
       ["Masking", "Dynamic masking at query time plus static masking during processing."],
       ["Sensitivity labels", "Information Protection labels applied to items; labels travel with "
        "the data."],
       ["Data residency", "Capacity and gateway virtual machines hosted in a Canadian region; no "
        "cross-geo processing."],
       ["Encryption", "Encrypted at rest by default; customer-managed keys evaluated if policy "
        "requires."],
       ["Export controls", "Publish-to-web disabled; export and download restricted on sensitive "
        "workspaces."]],
      widths=[1.7, 4.8])
para("Enforcement note: OneLake row/column security and dynamic data masking restrict "
     "non-privileged (Viewer) users. Workspace administrators and owners see unmasked data — "
     "which is precisely why administrative access is group-based, elevated just in time, and "
     "audited.", size=10, italic=True)

# ================================================================== GOVERNANCE
h("Governance", 1)
h("Operating model", 2)
table(["Role", "Owns", "Key responsibilities"],
      [["Platform Owner", "The Fabric platform and capacity",
        "Capacity, tenant and admin settings, standards, release management."],
       ["Capacity Admin", "Capacity sizing and guardrails",
        "Monitoring, throttling and surge guardrails, scale and rebalance path."],
       ["Platform Admins", "Tenant, gateways, releases",
        "Tenant settings, gateway upkeep, deployment pipelines, incident response."],
       ["Domain Owners", "Their business area's data",
        "Accountable for the data in their domain (Pensions, Employer, Member)."],
       ["Data Stewards", "Meaning, quality and access",
        "Definitions, quality rules and access decisions for their domain."],
       ["Data Product Owners", "Governed Gold products",
        "Each Gold output has an owner, a documented purpose and a consumer contract."]],
      widths=[1.5, 1.8, 3.2])

h("Tenant guardrails", 2)
table(["Tenant setting", "Recommended position"],
      [["Service principals can use Fabric APIs", "Enabled only for the BCPC-FAB-SPN-* group, "
        "not organisation-wide"],
       ["Publish to web", "Disabled"],
       ["External / guest sharing", "Disabled, or tightly restricted with approval"],
       ["Export and download", "Restricted for sensitive workspaces"],
       ["Sensitivity labels", "Enabled and enforced; mandatory labelling on new items"],
       ["Audit logging", "Enabled; logs shipped to BCPC's security monitoring"],
       ["Workspace creation", "Restricted to the platform team and a request process"],
       ["Capacity assignment", "Restricted — controls which workspaces run on which capacity"],
       ["Conditional Access", "Enforced on the Fabric and Power BI service"]],
      widths=[2.6, 3.9])

h("Optional tooling", 2)
para("The recommendation is to run on Fabric's built-in governance first — catalog, lineage, "
     "endorsement, data quality, audit, row/column security and sensitivity labels — and to add "
     "the tools below only where a specific requirement calls for them, since each carries "
     "licence and operational cost.", size=10.5)
table(["Tool", "What it adds", "Adopt when"],
      [["Microsoft Purview", "Enterprise catalog and data map, glossary at scale, DLP policies "
        "for Fabric, unified compliance reporting.",
        "Governance must span beyond Fabric, or a formal DLP programme is required."],
       ["Microsoft Sentinel", "SIEM — correlation, threat detection and response over Fabric and "
        "Entra audit logs.",
        "A security operations centre is mandated and logs must feed central detection."],
       ["Defender for Cloud", "Posture management for the surrounding Azure resources.",
        "Broader Azure posture management is in scope."],
       ["Customer-managed keys", "Bring-your-own encryption keys for OneLake and capacity.",
        "Policy requires key custody beyond Microsoft-managed encryption."]],
      widths=[1.4, 2.6, 2.5])
para("Neither Purview nor Sentinel is required for the platform to be secure and governed at "
     "launch. Purview is the most likely near-term addition given classification and DLP "
     "expectations for a crown corporation.", size=10, italic=True)

# ================================================================ CONNECTIVITY
h("Connectivity and gateway architecture", 1)
para("Two gateways are used, each for the role Microsoft supports it in. They are complementary, "
     "not competing.", size=10.5)
figure(os.path.join(OUT, "d4_gateway.png"),
       "Figure 4 — Private connectivity. Path A is the Microsoft-managed default; Path B is the "
       "self-hosted path, kept to the minimum necessary footprint.")

table(["Dimension", "VNet data gateway", "On-premises data gateway"],
      [["Hosting", "Microsoft-managed — nothing to install or patch",
        "BCPC-hosted Windows virtual machines, patched monthly"],
       ["Where it lives", "Injected into a delegated subnet of the BCPC virtual network",
        "A machine BCPC owns, on-premises or an Azure VM"],
       ["Reaches", "Azure and cloud sources over private endpoint or the virtual network",
        "On-premises and private sources; VNet resources when hosted on an Azure VM"],
       ["Private connectivity", "Private endpoints / Private Link — no public exposure",
        "Outbound-only; private over ExpressRoute or VPN"],
       ["Oracle mirroring", "Not supported", "Required — this is the only supported gateway"],
       ["High availability", "Managed by Microsoft", "BCPC clusters the gateway machines"],
       ["Capacity requirement", "F8 or above recommended — BCPC runs F64", "Any"],
       ["Operational cost", "Lower — managed hardware",
        "Virtual machine cost, patching, clustering and monitoring"]],
      widths=[1.5, 2.5, 2.5], size=9)

h("Which gateway a mirrored source needs", 2)
para("A gateway is not an Oracle-only concern: any mirroring source behind a firewall needs one. "
     "What differs is which gateway each source accepts.", size=10.5)
table(["Mirroring source", "On-premises gateway", "VNet gateway", "Note"],
      [["Oracle (on-prem, Azure VM, OCI, Exadata)", "Required", "Not supported",
        "Mandatory regardless of reachability"],
       ["SQL Server", "Yes", "Yes", "Either gateway when behind a firewall"],
       ["Snowflake", "Yes", "Yes", "Private Link to Fabric not yet supported"],
       ["Azure SQL Database / Managed Instance", "Yes", "Yes", "Either, when not public"],
       ["Azure Database for PostgreSQL", "—", "Yes", "VNet gateway for private servers"],
       ["Google BigQuery (preview)", "Yes", "—", "On-premises gateway"],
       ["Cosmos DB, Fabric SQL DB, open mirroring", "—", "—",
        "No gateway in the replication path"]],
      widths=[2.3, 1.3, 1.1, 1.8], size=9)

h("Position for BCPC", 2)
bullets([
    ("Short term", "reuse the existing VNet data gateway for Azure and private-endpoint sources "
     "and for pipeline ingestion; stand up a dedicated, highly available on-premises gateway only "
     "where it is required, principally Oracle mirroring."),
    ("Long term", "treat the VNet data gateway as the default managed path and minimise "
     "self-hosted gateway sprawl; as sources move to Azure or expose private endpoints, shift "
     "them across and shrink the on-premises footprint."),
    ("Throughout", "credentials in Key Vault, read-only source accounts, and governed control "
     "over who may create gateways and connections."),
], size=10.5)
para("Note: Oracle mirroring requires broad database grants on the source. That is a further "
     "reason to isolate the on-premises gateway and its credentials from every other connection.",
     size=10, italic=True)

# ============================================================= DECISIONS ETC.
h("Key decisions", 1)
para("These are the decisions TAG is being asked to endorse.", size=10.5)
table(["Area", "Decision"],
      [["Capacity and residency", "One Fabric F64 capacity hosted in a Canadian region, with "
        "gateway virtual machines in the same region."],
       ["Data architecture", "A medallion lakehouse on OneLake (Bronze → Silver → Gold). A "
        "warehouse is added only where a workload genuinely requires T-SQL write."],
       ["Configuration as code", "Sources, quality rules, models and security policies are held "
        "as configuration in git and promoted between environments."],
       ["Environment isolation", "Separate workspaces, Entra groups, service principals, key "
        "vaults and source connections for DEV, UAT and PROD."],
       ["Access model", "Access granted to Entra security groups only, with Conditional Access "
        "on every session and just-in-time elevation for privileged and production roles."],
       ["Connectivity", "The managed VNet data gateway is the default; a hardened, highly "
        "available on-premises gateway is used only where Microsoft requires it."],
       ["Governance tooling", "Begin with Fabric's built-in governance; Purview and Sentinel are "
        "optional, added against a specific requirement."],
       ["Ingestion pattern", "Framework ingestion into Bronze for batch sources; Fabric Mirroring "
        "where a source genuinely needs near-real-time replication."]],
      widths=[1.7, 4.8])

h("Assumptions", 1)
bullets([
    "The F64 capacity is available and can be assigned to the project workspaces.",
    "The existing VNet data gateway, already in use for Power BI, can be reused.",
    "BCPC identity services can provision the required Entra security groups and Conditional "
    "Access policies.",
    "The tenant setting permitting service principals to use Fabric APIs can be enabled, scoped "
    "to a named group rather than organisation-wide.",
    "The Oracle source can meet Microsoft's prerequisites for mirroring, including archive "
    "logging and supplemental logging, and a supported gateway version.",
    "Non-production environments will be supplied with masked or synthetic data; no unmasked "
    "member data is loaded outside production.",
    "Hosting in a Canadian region satisfies BCPC's data-residency position.",
    "PIA, STRA and AI-assessment requirements are confirmed through this submission before the "
    "platform moves beyond a pilot.",
], size=10.5)

h("Risks and mitigations", 1)
table(["Risk", "Rating", "Mitigation", "Owner"],
      [["Sensitive member data in a new platform", "High",
        "The pilot runs in an isolated environment on masked or synthetic data. The security and "
        "access model is reviewed with ITSO before any real data is loaded.", "ITE / ITSO"],
       ["Privileged access to production", "High",
        "No standing human change access to production; changes flow only through the deployment "
        "pipeline executed by a service principal, with administrator roles elevated just in time.",
        "Platform Owner"],
       ["Broad database grants required for Oracle mirroring", "Medium",
        "The on-premises gateway is dedicated, hardened and clustered, its credentials isolated "
        "in a per-environment key vault, and read-only accounts used everywhere else.",
        "Platform Owner"],
       ["Capacity contention and cost overrun", "Medium",
        "Production on dedicated capacity with surge protection and throttling guardrails; "
        "consumption monitored against the capacity allowance throughout the pilot.",
        "Capacity Admin"],
       ["Ungoverned workspace and report sprawl", "Medium",
        "Workspace creation restricted to a request-and-approve process; every workspace "
        "provisioned to standard from code; only reviewed Gold products certified.",
        "Data Governance"],
       ["Skills and adoption", "Low",
        "Adding a source is a configuration change rather than new code, keeping effort per "
        "source small; runbooks and working guides maintained alongside the platform.",
        "Project team"]],
      widths=[1.8, 0.7, 3.2, 0.9], size=9)

h("Out of scope for this phase", 1)
bullets([
    "Migrating production reporting off the current platform.",
    "Decommissioning any existing system.",
    "Loading unmasked member data into a non-production environment.",
    "Any public-internet data path to a BCPC source.",
    "Adopting Purview or Sentinel, unless a specific requirement is confirmed during the pilot.",
], size=10.5)

h("References", 1)
para("Microsoft documentation verified 2026-07-28.", size=10, italic=True)
for r in [
    "Fabric security overview — https://learn.microsoft.com/en-us/fabric/security/security-overview",
    "OneLake security (row and column level) — https://learn.microsoft.com/en-us/fabric/onelake/security/get-started-security",
    "Fabric service principals — https://learn.microsoft.com/en-us/fabric/admin/service-principal",
    "Fabric deployment pipelines — https://learn.microsoft.com/en-us/fabric/cicd/deployment-pipelines/intro-to-deployment-pipelines",
    "Information protection in Fabric — https://learn.microsoft.com/en-us/fabric/governance/information-protection",
    "Virtual network data gateway — https://learn.microsoft.com/en-us/data-integration/vnet/overview",
    "On-premises data gateway — https://learn.microsoft.com/en-us/data-integration/gateway/service-gateway-onprem",
    "Fabric Mirroring overview — https://learn.microsoft.com/en-us/fabric/mirroring/overview",
    "Oracle mirroring prerequisites — https://learn.microsoft.com/en-us/fabric/mirroring/oracle",
    "Fabric deployment patterns — https://learn.microsoft.com/en-us/azure/architecture/data-guide/technology-choices/fabric-deployment-patterns",
    "Conditional Access — https://learn.microsoft.com/en-us/entra/identity/conditional-access/overview",
    "Privileged Identity Management — https://learn.microsoft.com/en-us/entra/id-governance/privileged-identity-management/pim-configure",
]:
    para(r, size=9.5, space_after=2)

para("Diagram provenance: the figures are built from the official Microsoft Azure architecture "
     "icons and the Microsoft Fabric product icons. Oracle is a third-party product with no "
     "Microsoft icon and is drawn as a labelled element using a neutral database mark.",
     size=9, italic=True, color=GREY)

doc.save(DST)
print("saved:", DST)
