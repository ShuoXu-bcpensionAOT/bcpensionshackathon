"""Build WORKING_GUIDE.docx — a Word version of the working guide with Mermaid diagrams
rendered to images (via kroki.io) and the key reference tables. For readers who prefer docx.

    python control_plane/docs/build_docx.py
"""
import io
from pathlib import Path

import requests
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = Path(__file__).resolve().parent / "WORKING_GUIDE.docx"

# --- Mermaid diagram sources (kroki renders these to PNG) ---
DIAGRAMS = {
    "architecture": """flowchart LR
  S[(SQL Server source)]
  T["config_db (Fabric SQL DB)\\ndatasource, source_object, dq_rule,\\nmodel, gold_object, gold_dependency,\\nsteps, pbi_dataset"]
  MAIN[[cp_pl_main per load group]]
  PLAN[cp_plan planner]
  W[workers: metadata / bronze / silver / gold]
  B[(bronze)]
  SI[(silver)]
  G[(gold)]
  M[(metadata: state + logs)]
  T -->|pyodbc/AAD| PLAN
  MAIN --> PLAN --> W
  S -->|JDBC| B --> SI --> G
  W --> B & SI & G
  W -->|audit/logs| M
  T -.YAML export/promote.-> GIT[(git config/*.yml)]""",

    "medallion": """flowchart LR
  SRC[(Source table)] -->|extract full/incremental| BR[bronze: raw + control cols]
  BR -->|snake_case, dedupe by key, row-hash| SV[silver curated]
  SV -->|DQ rules| Q{pass?}
  Q -->|fail error rule| QT[quarantine_target]
  Q -->|pass| SVT[silver table]
  SVT -->|source-query sq_*| ST[stage_*]
  ST -->|SCD1 / SCD2 / fact merge| GD[gold: dim_* / fact_*]""",

    "orchestration": """flowchart TD
  P[PlanSteps: cp_plan steps] --> M{load_metadata active?}
  M -->|yes| MI[[cp_pl_metadata]] --> B
  M -->|no| B{load_bronze active?}
  B -->|yes| BI[[cp_pl_bronze]] --> S
  B -->|no| S{load_silver active?}
  S -->|yes| SI[[cp_pl_silver]] --> G
  S -->|no| G{load_gold active?}
  G -->|yes| GI[[cp_pl_gold]] --> R
  G -->|no| R{refresh_pbi active?}
  R -->|yes| RI[[cp_pl_pbi]] --> DONE([done])
  R -->|no| DONE
  BI -.on fail.-> F[cp_log_fail to pipeline_run_log, then Fail]
  F -.-> X([main fails])""",

    "planner_worker": """flowchart LR
  PL[cp_plan reads config_db via pyodbc] -->|exit JSON list| FE[ForEach item]
  FE --> W1[worker object 1]
  FE --> W2[worker object 2]
  FE --> W3[worker object N]
  W1 & W2 & W3 --> LG[(lakehouse)]""",

    "erd": """erDiagram
  datasource   ||--o{ source_object    : has
  source_object ||--o{ dq_rule          : validated_by
  model        ||--o{ gold_object       : contains
  gold_object  ||--o{ gold_dependency   : parent_or_child
  datasource      { int source_id }
  source_object   { string object_id }
  dq_rule         { string rule_id }
  model           { int model_id }
  gold_object     { string gold_object_id }
  gold_dependency { string parent_child }
  steps           { string step_key }
  pbi_dataset     { string dataset_id }""",

    "gold_dag": """flowchart TD
  DC[dim_category] --> DSC[dim_subcategory] --> DP[dim_product]
  DT[dim_territory] --> DCU[dim_customer]
  DP --> F[fact_sales_order]
  DCU --> F
  DT --> F
  F --> FT[fact_sales_by_territory]""",

    "promotion": """flowchart LR
  DEV[DEV config_db: edit via T-SQL] -->|cp_export_config| YML[config/*.yml git]
  YML -->|PR / merge| MAIN[(main branch)]
  MAIN -->|cp_config| UAT[UAT config_db]
  MAIN -->|cp_config| PROD[PROD config_db]""",

    "deployment": """flowchart LR
  GH[git repo] --> CI[Deploy Control Plane: GitHub Action / cp_bootstrap]
  CI -->|az login SP| WS[workspace base-env]
  CI --> LHK[lakehouses + config_db]
  CI --> VL[cp_vars]
  CI --> NB[notebooks]
  CI --> PL[data pipelines]
  CI --> CFG[load config-as-code]""",
}


def render(mermaid):
    r = requests.post("https://kroki.io/mermaid/png",
                      json={"diagram_source": mermaid}, timeout=40)
    r.raise_for_status()
    return io.BytesIO(r.content)


def main():
    imgs = {k: render(v) for k, v in DIAGRAMS.items()}
    print("rendered", len(imgs), "diagrams")
    d = Document()

    def h(t, lvl):
        d.add_heading(t, level=lvl)

    def p(t):
        d.add_paragraph(t)

    def bullets(items):
        for i in items:
            d.add_paragraph(i, style="List Bullet")

    def fig(key, caption):
        d.add_picture(imgs[key], width=Inches(6.2))
        d.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        c = d.add_paragraph(caption)
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.runs[0].italic = True
        c.runs[0].font.size = Pt(9)

    def table(headers, rows):
        t = d.add_table(rows=1, cols=len(headers))
        t.style = "Light Grid Accent 1"
        for i, hh in enumerate(headers):
            cell = t.rows[0].cells[i]
            cell.text = hh
            cell.paragraphs[0].runs[0].bold = True
        for row in rows:
            cells = t.add_row().cells
            for i, v in enumerate(row):
                cells[i].text = str(v)
        d.add_paragraph()

    # Title
    title = d.add_heading("Fabric Control Plane — Working Guide", level=0)
    sub = d.add_paragraph("Metadata-driven, config-first lakehouse platform for Microsoft Fabric")
    sub.runs[0].italic = True

    h("1. Architecture at a glance", 1)
    p("You describe WHAT to load and build in config tables; the framework does the HOW — "
      "extract to bronze, curate to silver (with data quality), build a gold star schema — "
      "orchestrated by Fabric Data Pipelines and promotable across environments from git.")
    fig("architecture", "Figure 1. End-to-end architecture")
    fig("medallion", "Figure 2. Medallion data flow (bronze → silver → gold)")
    h("Key principles", 2)
    bullets([
        "Config is data, not code: authored config lives in config_db (T-SQL editable); "
        "runtime state/logs live in the metadata lakehouse. Only authored config is promoted.",
        "Items vs data promotion: Fabric items promote together via deployment; config data "
        "promotes as YAML (cp_export_config → git → cp_config). Runtime state never promotes.",
        "Notebooks are param-driven workers: they never read config; a planner (cp_plan) reads "
        "config and the pipeline fans work out to the workers.",
        "Zero hardcoding: workspace + lakehouse IDs resolve at runtime; environment values come "
        "from the cp_vars Variable Library.",
    ])

    h("2. Resources and what they do", 1)
    h("2.1 Lakehouses", 2)
    table(["Lakehouse", "Purpose", "Contents"], [
        ["bronze", "Raw landing (as-extracted)", "<source>_<schema>_<table> Delta + control cols"],
        ["silver", "Curated, deduped, DQ-passed", "same names; snake_case, _row_hash; quarantine_<target>"],
        ["gold", "Star schema (business)", "dim_* / fact_*; plus stage_*"],
        ["metadata", "Runtime state + logs + config mirror", "audit/state tables; error tracebacks in Files"],
    ])
    h("2.2 config_db (Fabric SQL Database)", 2)
    p("Holds the authored config tables (section 4). Edited with T-SQL. Auto-mirrors to OneLake; "
      "tooling reads/writes via pyodbc + Entra token.")
    h("2.3 Variable Library cp_vars", 2)
    p("Per-environment values, swappable via value sets: lakehouse names, source_server, "
      "source_connection. Notebooks read it at runtime.")
    h("2.4 Notebooks", 2)
    table(["Notebook", "Folder", "Role", "Key parameters"], [
        ["cp_framework", "utility", "Shared library (%run by all): paths, config read, JDBC, Delta helpers, gold writers, DAG sort, audit", "—"],
        ["cp_plan", "utility", "Reads config_db, returns the ForEach work-list", "load_group, plan_type"],
        ["cp_log_fail", "utility", "Writes a failure row to pipeline_run_log", "pipeline_name, run_id, load_group, activity, message"],
        ["metadata_worker", "notebook", "Start run; discover source schema; log drift", "run_id, load_group, src_user, src_password"],
        ["bronze_worker", "notebook", "Extract ONE object to bronze (full/incremental)", "run_id, object_json, src_user, src_password"],
        ["silver_worker", "notebook", "Build silver for ONE object: dedupe, DQ→quarantine", "run_id, object_json"],
        ["gold_runner", "notebook", "Build ONE model's gold objects in DAG order (runs sq_*)", "run_id, model_id"],
        ["sq_*", "sourcequery", "Source-query builder per gold object: silver→stage→gold", "run_id"],
    ])
    h("2.5 Pipelines (in the 'pipeline' folder)", 2)
    table(["Pipeline", "Purpose", "Parameters"], [
        ["cp_pl_main", "Orchestrator per load group: read steps, run 5 steps in order, skip inactive, fail-fast", "load_group, run_id, src_user, src_password"],
        ["cp_pl_metadata", "Runs metadata_worker", "load_group, run_id, src_user, src_password"],
        ["cp_pl_bronze", "cp_plan(objects) → ForEach → bronze_worker", "load_group, run_id, src_user, src_password"],
        ["cp_pl_silver", "cp_plan(objects) → ForEach → silver_worker", "load_group, run_id"],
        ["cp_pl_gold", "cp_plan(models) → ForEach → gold_runner", "load_group, run_id"],
        ["cp_pl_pbi", "cp_plan(datasets) → ForEach → Power BI REST refresh (scaffold)", "load_group, run_id"],
    ])
    fig("orchestration", "Figure 3. Main orchestrator: sequential, is_active-gated, fail-fast")
    fig("planner_worker", "Figure 4. Planner–worker pattern inside a child pipeline")
    p("Every pipeline: on a work-activity failure → cp_log_fail writes to pipeline_run_log, then a "
      "Fail activity re-fails so cp_pl_main errors out (fail-fast).")

    h("3. How to operate", 1)
    bullets([
        "Deploy an environment: python control_plane/deploy/cp_bootstrap.py <workspace_base> <ENV> "
        "(idempotent, deploy-only), or run the Deploy Control Plane GitHub Action.",
        "Run the pipeline: trigger cp_pl_main(load_group, run_id, src_user, src_password).",
        "Trial-capacity note: don't run two full pipelines on one capacity at once — Spark session "
        "slots are limited; run environments sequentially.",
        "Author/change config: edit config_db via T-SQL, then cp_export_config (SQL→YAML, commit) "
        "and cp_config (YAML→target SQL) to promote.",
        "Add a source object → source_object row; Add a gold table → sq_<name> + gold_object (+ "
        "gold_dependency); Add a DQ rule → dq_rule row; Toggle a step → steps.is_active.",
    ])

    h("4. config_db table reference (authored config)", 1)
    p("is_active is a BIT on every authored table; only active rows are processed. "
      "Standard/enumerated values are noted below.")
    fig("erd", "Figure 5. Config database — entity relationships")
    h("4.1 datasource — source systems", 2)
    table(["Column", "Notes / allowed values"], [
        ["source_id (PK)", "unique id"], ["source_name", "display name"],
        ["source_type", "SQL"], ["database_name", "source database"],
        ["load_group", "the run unit; cp_pl_main runs one load group"],
        ["ingestion_mode", "custom_jdbc"], ["is_active", "BIT"],
    ])
    h("4.2 source_object — objects to ingest", 2)
    table(["Column", "Notes / allowed values"], [
        ["object_id (PK)", "stable logical id (e.g. customer)"],
        ["source_id (FK)", "→ datasource"],
        ["source_schema, source_table", "source location"],
        ["target_name", "bronze/silver Delta table name"],
        ["load_type", "full = overwrite each run · incremental = append rows past the watermark"],
        ["key_columns_json", 'business key(s), JSON array e.g. ["SalesOrderID","SalesOrderDetailID"]'],
        ["watermark_column", "column for incremental (e.g. ModifiedDate)"],
        ["watermark_type", "datetime"],
        ["processing_state", "ACTIVE (only ACTIVE objects run)"],
        ["is_active", "BIT"],
    ])
    p("Incremental: first run pulls all and records max(watermark); later runs pull rows past the "
      "stored watermark and append. Full loads overwrite. Complex source types (xml/geography/"
      "geometry/hierarchyid/varbinary/image/sql_variant) are excluded automatically.")
    h("4.3 dq_rule — data-quality rules (evaluated on silver)", 2)
    p("column_name is the SILVER column (snake_case), e.g. total_due, customer_id.")
    table(["rule_type", "Fill these columns", "Passes when", "Example"], [
        ["not_null", "column_name", "column IS NOT NULL", "customer_id not null"],
        ["min", "column_name, min_value", "column ≥ min (NULL passes)", "total_due ≥ 0"],
        ["max", "column_name, max_value", "column ≤ max (NULL passes)", "discount ≤ 1"],
        ["allowed_values", "column_name, allowed_values_json", "column ∈ list (NULL passes)", 'person_type ∈ ["EM","IN"]'],
        ["expression", "rule_expression (column_name optional)", "the boolean expression is TRUE", "order_qty > 0 AND unit_price >= 0"],
    ])
    p("severity: error = failing rows quarantined and excluded from silver (written to "
      "quarantine_<target>); warn = counted in dq_result only. All rule counts are recorded in "
      "dq_result.")
    p('Example (T-SQL):  INSERT INTO dbo.dq_rule (rule_id, object_id, column_name, rule_type, '
      "min_value, severity, is_active) VALUES "
      "('soh_total_due_nonneg','sales_order_header','total_due','min',0,'error',1);")
    h("4.4 model — gold data models", 2)
    table(["Column", "Notes"], [["model_id (PK)", ""], ["model_name", "e.g. sales_star"],
                                 ["load_group", "gold-phase load group"], ["is_active", "BIT"]])
    h("4.5 gold_object — gold tables", 2)
    table(["Column", "Notes / allowed values"], [
        ["gold_object_id (PK)", "e.g. dim_customer"], ["model_id (FK)", "→ model"],
        ["gold_type", "scd1 (upsert by key) · scd2 (history: _is_current, _effective_start/end_ts, _row_hash) · fact (upsert by key)"],
        ["stage_table", "staging table name (stage_* in gold)"],
        ["gold_table", "final gold table name"],
        ["business_key_columns_json", 'key(s), snake_case, e.g. ["product_key"]'],
        ["source_query_notebook", "the sq_* notebook that builds this object's stage"],
        ["is_active", "BIT"],
    ])
    h("4.6 gold_dependency — build order (DAG)", 2)
    table(["Column", "Notes"], [["parent_gold_object_id", "built first"],
                                ["child_gold_object_id", "built after parent"]])
    fig("gold_dag", "Figure 6. Example gold DAG (sales_star) — order gold_runner derives")
    h("4.7 steps — orchestration steps per load group", 2)
    table(["Column", "Notes / allowed values"], [
        ["load_group", "INT"], ["step_order", "1..5 (execution order)"],
        ["step_key", "load_metadata · load_bronze · load_silver · load_gold · refresh_pbi"],
        ["child_pipeline", "pipeline to invoke"], ["is_active", "BIT — inactive steps are skipped"],
    ])
    h("4.8 pbi_dataset — Power BI refresh targets (scaffold)", 2)
    table(["Column", "Notes"], [["dataset_id (PK), workspace_id, dataset_name, load_group, is_active",
                                 "REST refresh issued per active row"]])

    h("5. Runtime state / log tables (generated — never authored, never promoted)", 1)
    table(["Table", "What it records"], [
        ["ingestion_run", "one row per run start/finish"],
        ["object_load_run", "per object × layer: status + source/target/quarantine counts"],
        ["watermark_state", "latest watermark per object (incremental)"],
        ["source_column", "discovered source schema snapshot (drift baseline)"],
        ["schema_drift_event", "column added/removed vs previous snapshot"],
        ["dq_result", "per rule: passed/failed counts, PASS/FAIL"],
        ["quarantine_<target>", "(silver) rows that failed an error-severity DQ rule"],
        ["pipeline_run_log", "pipeline failures (pipeline, activity, error message)"],
    ])

    h("6. Promotion & environments", 1)
    fig("promotion", "Figure 7. Config authoring & promotion loop (tables are source of truth)")
    fig("deployment", "Figure 8. Item deployment per environment")
    bullets([
        "Items deploy per environment via the bootstrap or CI/CD; naming is <workspace_base>-<environment>.",
        "Environment values come from cp_vars value sets; deploy parameters from environments/<env>.yml.",
        "Config data promotes as YAML: cp_export_config (SQL→YAML) then cp_config (YAML→target SQL).",
        "Secrets: source password passed at run time; SP client id/secret drive CI auth. Key Vault + "
        "native pipeline SQL Lookups are planned upgrades once an SP exists.",
    ])

    d.save(str(OUT))
    print("wrote WORKING_GUIDE.docx")


if __name__ == "__main__":
    main()
