"""Build CICD.docx — Word version of the CI/CD guide (GitHub + Azure DevOps streams),
with diagrams rendered white-background via kroki. Run: python control_plane/docs/build_cicd_docx.py
"""
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from build_docx import render  # white-bg mermaid renderer

OUT = Path(__file__).resolve().parent / "CICD.docx"

DIAGRAMS = {
    "github": """flowchart LR
  P[git push] --> GH[GitHub Actions: Deploy Control Plane]
  GH -->|OIDC federated login| SP[Entra service principal]
  GH -->|GitHub Environment| A{approved?}
  A -->|yes| BOOT[cp_bootstrap.py base env]
  BOOT --> WS[workspace base-env: items + config]""",
    "devops": """flowchart LR
  R[Azure Repos / GitHub] --> ADO[Azure DevOps Pipeline]
  ADO -->|WIF service connection| SP[Entra service principal]
  ADO -->|variable group linked to| KV[(Key Vault)]
  ADO -->|Environment approval| A{approved?}
  A -->|yes| BOOT[cp_bootstrap.py base env]
  BOOT --> WS[workspace base-env]""",
    "multienv": """flowchart LR
  ENVF[environments/dev,uat,prod.yml] --> BOOT[cp_bootstrap.py]
  SEC[env-scoped secrets] --> BOOT
  BOOT --> D[base-DEV]
  BOOT --> U[base-UAT]
  BOOT --> PR[base-PROD]""",
}


def main():
    imgs = {k: render(v) for k, v in DIAGRAMS.items()}
    d = Document()

    def h(t, l):
        d.add_heading(t, level=l)

    def p(t):
        d.add_paragraph(t)

    def bullets(items):
        for i in items:
            d.add_paragraph(i, style="List Bullet")

    def fig(key, cap):
        d.add_picture(imgs[key], width=Inches(6.2))
        d.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        c = d.add_paragraph(cap); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.runs[0].italic = True; c.runs[0].font.size = Pt(9)

    def table(headers, rows):
        t = d.add_table(rows=1, cols=len(headers)); t.style = "Light Grid Accent 1"
        for i, hh in enumerate(headers):
            t.rows[0].cells[i].text = hh; t.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        for row in rows:
            cs = t.add_row().cells
            for i, v in enumerate(row):
                cs[i].text = str(v)
        d.add_paragraph()

    d.add_heading("CI/CD Guide — Deploying the Fabric Control Plane", 0)
    s = d.add_paragraph("Two streams: GitHub Actions and Azure DevOps Pipelines. Both call the "
                        "same cp_bootstrap.py (idempotent, deploy-only).")
    s.runs[0].italic = True

    h("Two kinds of secret", 1)
    table(["Context", "What", "Where it belongs"], [
        ["Deploy-time", "Identity that provisions/deploys Fabric items",
         "CI identity — Entra SP via OIDC / Workload Identity Federation (preferred) or client secret in CI store"],
        ["Run-time", "Source DB password used when cp_pl_main runs",
         "Azure Key Vault, read by the notebook at run time — not needed for deploy"],
    ])
    p("The CI/CD pipelines only DEPLOY (provision items + load config). They do not need the "
      "source password; that is a run-time secret (see Vault).")

    h("Service principal (shared prerequisite)", 1)
    bullets([
        "Enable tenant settings: service principals can use Fabric APIs, and can create workspaces/connections/pipelines.",
        "Add the SP as Admin on each target workspace.",
        "Grant capacity assignment/contributor rights.",
        "Grant config_db access: CREATE USER [<sp>] FROM EXTERNAL PROVIDER; ALTER ROLE db_owner ADD MEMBER [<sp>];",
        "Prefer federated credentials (OIDC) over a client secret — nothing to store or rotate.",
    ])

    h("Repository layout (both platforms)", 1)
    p(".github/workflows/deploy.yml (GitHub)  ·  control_plane/cicd/azure-pipelines.yml (DevOps)  ·  "
      "control_plane/environments/{dev,uat,prod}.yml  ·  config/ variable_library/ notebooks/ deploy/ requirements.txt")

    h("Stream A — GitHub Actions", 1)
    fig("github", "Figure 1. GitHub Actions deploy flow")
    bullets([
        "Push the repo; the workflow is .github/workflows/deploy.yml.",
        "Settings → Environments: create DEV/UAT/PROD; add required reviewers on UAT/PROD; scope secrets per environment.",
        "Secure auth (recommended): OIDC — add a federated credential on the Entra app "
        "(issuer token.actions.githubusercontent.com, subject repo:<org>/<repo>:environment:<ENV>), store only AZURE_CLIENT_ID.",
        "Fallback: AZURE_CLIENT_ID + AZURE_CLIENT_SECRET secrets; workflow runs az login --service-principal.",
        "Run: Actions → Deploy Control Plane → Run workflow → pick environment.",
    ])

    h("Stream B — Azure DevOps Pipelines", 1)
    fig("devops", "Figure 2. Azure DevOps deploy flow")
    bullets([
        "Push to Azure Repos (or connect a GitHub repo). New pipeline → Existing YAML → control_plane/cicd/azure-pipelines.yml.",
        "Secure auth (recommended): ARM service connection 'fabric-sp' with Workload Identity Federation — no stored secret.",
        "Secrets from Key Vault: Library → Variable group 'fabric-control-plane' → link secrets from Azure Key Vault; grant the service connection Get/List.",
        "Environments: create fabric-dev/uat/prod; add Approvals & checks on uat/prod.",
        "Multi-stage promotion: chain DEV → UAT → PROD deployment jobs; environment approvals are the gates.",
    ])

    h("Multi-environment model", 1)
    fig("multienv", "Figure 3. One codebase → per-environment workspaces")
    bullets([
        "One workspace per environment: <workspace_base>-<ENV>.",
        "Per-env inputs from environments/<env>.yml (non-secret, in git).",
        "Per-env runtime values from cp_vars value sets (lakehouse names, source_server).",
        "Per-env secrets scoped by the platform (GitHub Environment secrets / DevOps KV-linked variable groups).",
    ])

    h("Vault & secure storage", 1)
    bullets([
        "Deploy-time identity: prefer OIDC / Workload Identity Federation (no stored secret); else client secret in a GitHub Environment secret or KV-linked DevOps variable group; rotate.",
        "Run-time source secret: store in Azure Key Vault; the worker reads it via "
        "notebookutils.credentials.getSecret('https://<vault>.vault.azure.net/','source-db-password'); grant the workspace identity Key Vault Secrets User.",
        "Never in git: SP secrets, DB passwords, PATs, connection strings — only non-secret params and config structure.",
    ])

    h("GitHub vs Azure DevOps — at a glance", 1)
    table(["Concern", "GitHub Actions", "Azure DevOps"], [
        ["CI file", ".github/workflows/deploy.yml", "control_plane/cicd/azure-pipelines.yml"],
        ["Secure auth", "OIDC federated credential (or SP secret)", "WIF service connection (or SP secret)"],
        ["Secrets store", "Repo/Environment Secrets", "Variable group linked to Key Vault"],
        ["Approvals", "Environments → required reviewers", "Environments → approvals & checks"],
        ["Multi-env", "env input → GitHub Environment", "parameter → Environment; multi-stage promotion"],
        ["Vault", "Key Vault via az keyvault / OIDC", "Native KV-linked variable group"],
    ])

    d.save(str(OUT))
    print("wrote CICD.docx")


if __name__ == "__main__":
    main()
