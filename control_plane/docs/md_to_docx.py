"""Generic Markdown -> .docx converter (no pandoc needed). Handles headings, paragraphs with
**bold** / `inline code` / [links](url), bullet lists, fenced code blocks, tables, and rules.

    python md_to_docx.py GOVERNANCE_SECURITY.md   # -> GOVERNANCE_SECURITY.docx
"""
import base64
import io
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

MONO = "Consolas"
LINK = RGBColor(0x1A, 0x5F, 0xB4)
CODE = RGBColor(0xA3, 0x1F, 0x34)
MERMAID_INIT = "%%{init: {'theme':'neutral'}}%%\n"
_INLINE = re.compile(r"(\*\*.+?\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))")


def render_mermaid(src):
    """Render mermaid text -> PNG bytes flattened onto white. Prefer kroki (PNG),
    fall back to mermaid.ink. Returns None if neither is reachable."""
    import requests
    from PIL import Image
    body = MERMAID_INIT + src
    raw = None
    try:
        r = requests.post("https://kroki.io/mermaid/png", json={"diagram_source": body}, timeout=40)
        r.raise_for_status()
        raw = r.content
    except Exception as e:
        print(f"  kroki unavailable ({e}); using mermaid.ink")
        try:
            b = base64.urlsafe_b64encode(body.encode()).decode()
            r = requests.get("https://mermaid.ink/img/" + b, timeout=40)
            r.raise_for_status()
            raw = r.content
        except Exception as e2:
            print(f"  mermaid.ink unavailable ({e2}); leaving diagram as text")
            return None
    img = Image.open(io.BytesIO(raw)).convert("RGBA")
    bg = Image.new("RGBA", img.size, (255, 255, 255, 255))
    flat = Image.alpha_composite(bg, img).convert("RGB")
    buf = io.BytesIO()
    flat.save(buf, "PNG")
    buf.seek(0)
    return buf


def _runs(paragraph, text):
    """Add text to a paragraph, rendering **bold**, `code`, and [links](url)."""
    for part in _INLINE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            paragraph.add_run(part[2:-2]).bold = True
        elif part.startswith("`") and part.endswith("`"):
            r = paragraph.add_run(part[1:-1]); r.font.name = MONO; r.font.color.rgb = CODE
        elif part.startswith("["):
            m = re.match(r"\[([^\]]+)\]\(([^)]+)\)", part)
            r = paragraph.add_run(m.group(1)); r.font.color.rgb = LINK; r.font.underline = True
        else:
            paragraph.add_run(part)


def _cell(cell, text):
    cell.paragraphs[0].text = ""
    _runs(cell.paragraphs[0], text.strip())
    for r in cell.paragraphs[0].runs:
        r.font.size = Pt(9)


def convert(md_path):
    md = Path(md_path)
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)
    lines = md.read_text(encoding="utf-8").splitlines()
    i = 0
    while i < len(lines):
        ln = lines[i]
        # fenced code block (```lang ... ```)
        if ln.lstrip().startswith("```"):
            lang = ln.lstrip()[3:].strip().lower()
            i += 1
            buf = []
            while i < len(lines) and not lines[i].lstrip().startswith("```"):
                buf.append(lines[i]); i += 1
            i += 1
            body = "\n".join(buf)
            if lang == "mermaid":
                png = render_mermaid(body)
                if png is not None:
                    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.add_run().add_picture(png, width=Inches(6.2))
                    continue
                # fall through to text if rendering failed
            p = doc.add_paragraph()
            r = p.add_run(body)
            r.font.name = MONO; r.font.size = Pt(8.5)
            p.paragraph_format.left_indent = Pt(12)
            continue
        # table (line with | and next line is a separator)
        if ln.strip().startswith("|") and i + 1 < len(lines) and re.match(r"^\s*\|[\s:|-]+\|\s*$", lines[i + 1]):
            header = [c.strip() for c in ln.strip().strip("|").split("|")]
            i += 2
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
                i += 1
            t = doc.add_table(rows=1, cols=len(header)); t.style = "Light Grid Accent 1"
            for j, h in enumerate(header):
                _cell(t.rows[0].cells[j], h)
                for run in t.rows[0].cells[j].paragraphs[0].runs:
                    run.bold = True
            for row in rows:
                cells = t.add_row().cells
                for j in range(len(header)):
                    _cell(cells[j], row[j] if j < len(row) else "")
            doc.add_paragraph()
            continue
        # heading
        m = re.match(r"^(#{1,5})\s+(.*)$", ln)
        if m:
            lvl = len(m.group(1))
            h = doc.add_heading(level=min(lvl, 4))
            h.text = ""
            _runs(h, re.sub(r"[`*]", "", m.group(2)))
            i += 1
            continue
        # horizontal rule
        if re.match(r"^\s*---+\s*$", ln):
            i += 1
            continue
        # bullet
        m = re.match(r"^(\s*)[-*]\s+(.*)$", ln)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Pt(18 + len(m.group(1)) // 2 * 12)
            _runs(p, m.group(2))
            i += 1
            continue
        # blank / paragraph
        if ln.strip():
            _runs(doc.add_paragraph(), ln)
        i += 1

    out = md.with_suffix(".docx")
    doc.save(out)
    print(f"wrote {out.name}")


if __name__ == "__main__":
    convert(sys.argv[1] if len(sys.argv) > 1 else "GOVERNANCE_SECURITY.md")
